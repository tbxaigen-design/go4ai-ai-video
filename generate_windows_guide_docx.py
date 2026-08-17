import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_callout_box(doc, text_list, title="💡 LƯU Ý QUAN TRỌNG", bg_hex="F0F9FF", border_color="0284C7"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=220)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:top w:val="none"/>
            <w:right w:val="none"/>
            <w:bottom w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(title)
    run_t.bold = True
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(15, 23, 42)
    
    for line in text_list:
        p_line = cell.add_paragraph()
        p_line.paragraph_format.space_before = Pt(2)
        p_line.paragraph_format.space_after = Pt(3)
        run = p_line.add_run(line)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 41, 59)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def build_windows_docx():
    doc = Document()
    
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("💻 GO4AI AI VIDEO STUDIO")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(30, 64, 175) # #1E40AF
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("Cẩm Nang Hướng Dẫn Cài Đặt & Sử Dụng Chi Tiết Cho Máy Tính Windows 10 & 11\n(Dành cho người mới bắt đầu / No-code • Tự động hóa hoàn toàn)")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(71, 85, 105)
    run_sub.italic = True
    
    # Metadata Badge Table
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cols = [
        ("💻 Nền tảng", "Windows 10 / Windows 11 (64-bit)"),
        ("⚡ Phiên bản", "Beta 1.0 (Local Render)"),
        ("🤝 Hỗ trợ", "hocvien@go4ai.life")
    ]
    for i, (k, v) in enumerate(cols):
        cell = meta_table.cell(0, i)
        cell.width = Inches(2.15)
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{k}\n")
        r1.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = RGBColor(30, 41, 59)
        r2 = p.add_run(v)
        r2.font.name = 'Arial'
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # --- PHẦN 1: GIỚI THIỆU ---
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("1. Giới Thiệu & Điểm Khác Biệt")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "GO4AI AI Video Studio là giải pháp tạo video tự động bằng trí tuệ nhân tạo. "
        "Ứng dụng hoạt động theo cơ chế Render cục bộ (Local Render) trên máy tính Windows, "
        "giúp video của bạn được xử lý bảo mật tuyệt đối, xuất chuẩn MP4 Full HD cực nét mà không phụ thuộc vào hàng chờ trên đám mây."
    ).font.name = 'Arial'
    
    add_callout_box(
        doc,
        [
            "• Bạn KHÔNG cần cài đặt môi trường phức tạp trước, hệ thống có tính năng tự động tải Node.js portable nếu máy bạn chưa có.",
            "• Bạn KHÔNG cần gõ dòng lệnh nào — chỉ cần mở file .bat.",
            "• Khi mở file, nếu thấy một cửa sổ màu đen xuất hiện (Command Prompt), đó là ứng dụng đang tự động chạy ngầm."
        ],
        title="🌟 DÀNH CHO BẠN (NGƯỜI DÙNG NO-CODE)",
        bg_hex="ECFDF5",
        border_color="059669"
    )
    
    # --- PHẦN 2: CẤU HÌNH YÊU CẦU ---
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    r = h2.add_run("2. Cấu Hình Máy Tính Khuyến Nghị")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    config_table = doc.add_table(rows=5, cols=2)
    config_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_data = [
        ("Hệ điều hành", "Windows 10 (64-bit) hoặc Windows 11"),
        ("Bộ vi xử lý (CPU)", "Intel Core i3/i5 Gen 8+ hoặc AMD Ryzen 3 trở lên"),
        ("Bộ nhớ RAM", "Tối thiểu 8 GB RAM (Khuyến nghị 16 GB để render nhanh nhất)"),
        ("Ổ đĩa trống", "Tối thiểu 5 GB SSD trống"),
        ("Kết nối Internet", "Cần kết nối mạng trong lần đầu mở để tải tài nguyên và tạo giọng đọc AI")
    ]
    for row_idx, (k, v) in enumerate(c_data):
        cell_k = config_table.cell(row_idx, 0)
        cell_v = config_table.cell(row_idx, 1)
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        set_cell_background(cell_k, "F8FAFC")
        set_cell_background(cell_v, "FFFFFF")
        set_cell_margins(cell_k, 70, 70, 100, 100)
        set_cell_margins(cell_v, 70, 70, 100, 100)
        
        p1 = cell_k.paragraphs[0]
        r1 = p1.add_run(k)
        r1.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(9)
        
        p2 = cell_v.paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.name = 'Arial'
        r2.font.size = Pt(9)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # --- PHẦN 3: CÁC BƯỚC CÀI ĐẶT ---
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    r = h3.add_run("3. Hướng Dẫn Cài Đặt & Mở Ứng Dụng (3 Bước Đơn Giản)")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    # Bước 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("📥 BƯỚC 1: Tải bộ phần mềm từ GitHub")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(2, 132, 199)
    
    p_desc1 = doc.add_paragraph()
    p_desc1.paragraph_format.space_after = Pt(4)
    p_desc1.add_run(
        "1. Truy cập đường link tải: https://github.com/tbxaigen-design/go4ai-ai-video\n"
        "2. Nhìn sang góc phải, bấm vào nút màu xanh lá cây có chữ '< > Code'.\n"
        "3. Trong bảng hiện ra, bấm vào dòng 'Download ZIP'.\n"
        "4. File nén 'go4ai-ai-video-main.zip' sẽ tải về thư mục Downloads.\n"
        "5. Click chuột phải vào file .zip vừa tải ➔ Chọn 'Extract All...' (Giải nén tất cả) ➔ Chọn thư mục lưu (ví dụ Desktop hoặc ổ D:\\) ➔ Bấm Extract."
    ).font.name = 'Arial'
    
    # Bước 2
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("🚀 BƯỚC 2: Khởi động file 1-Chay-Tren-Windows.bat")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(2, 132, 199)
    
    p_desc2 = doc.add_paragraph()
    p_desc2.paragraph_format.space_after = Pt(4)
    p_desc2.add_run(
        "Mở thư mục vừa giải nén, bạn tìm file có biểu tượng bánh răng tên: 1-Chay-Tren-Windows.bat\n"
        "1. Nhấp đúp chuột (Double-click) vào file '1-Chay-Tren-Windows.bat'.\n"
        "2. Nếu màn hình Windows hiện bảng bảo vệ màu xanh 'Windows protected your PC' (SmartScreen):\n"
        "   👉 Bấm vào chữ 'More info' (Thông tin khác)\n"
        "   👉 Bấm tiếp nút 'Run anyway' (Vẫn chạy)\n"
        "3. Một cửa sổ đen (Command Prompt) sẽ xuất hiện và tự động cài đặt mọi thứ trong 2–4 phút."
    ).font.name = 'Arial'
    
    add_callout_box(
        doc,
        [
            "• Hệ thống tự động kiểm tra Node.js, cài thư viện render Remotion, tải bộ mã hoá video FFmpeg và giọng đọc AI.",
            "• Quá trình chỉ diễn ra 1 lần duy nhất trong lần đầu mở.",
            "• Vui lòng GIỮ NGUYÊN cửa sổ màu đen và không tắt trong lúc đang tải."
        ],
        title="⏳ QUÁ TRÌNH TỰ ĐỘNG CHUẨN BỊ (2-4 PHÚT)",
        bg_hex="FFFBEB",
        border_color="D97706"
    )
    
    # Bước 3
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("🌐 BƯỚC 3: Trải nghiệm GO4AI Studio trên trình duyệt")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(2, 132, 199)
    
    p_desc3 = doc.add_paragraph()
    p_desc3.paragraph_format.space_after = Pt(6)
    p_desc3.add_run(
        "Sau khi chuẩn bị xong, trình duyệt mặc định (Chrome / Edge / Cốc Cốc) sẽ tự động mở trang web: http://127.0.0.1:3075\n"
        "Giao diện Studio đã sẵn sàng để bạn tạo video AI đầu tiên!"
    ).font.name = 'Arial'
    
    # --- PHẦN 4: SỬ DỤNG HÀNG NGÀY ---
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(6)
    r = h4.add_run("4. Sử Dụng Hàng Ngày & Tắt Ứng Dụng")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    usage_table = doc.add_table(rows=2, cols=2)
    usage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    u_data = [
        ("🟢 Cách Mở App Hàng Ngày", "Chỉ cần nhấp đúp vào file '1-Chay-Tren-Windows.bat'. App sẽ mở ngay lập tức trong 2-3 giây. Bạn có thể Click chuột phải ➔ 'Send to' ➔ 'Desktop (create shortcut)' để mở nhanh từ màn hình chính."),
        ("🔴 Cách Tắt App Đúng Cách", "Khi không dùng nữa, bạn quay lại cửa sổ màu đen và bấm phím Ctrl + C hoặc bấm dấu [X] ở góc phải cửa sổ đen để đóng ứng dụng.")
    ]
    for row_idx, (title, desc) in enumerate(u_data):
        cell_t = usage_table.cell(row_idx, 0)
        cell_d = usage_table.cell(row_idx, 1)
        cell_t.width = Inches(2.3)
        cell_d.width = Inches(4.2)
        set_cell_background(cell_t, "F8FAFC")
        set_cell_background(cell_d, "FFFFFF")
        set_cell_margins(cell_t, 100, 100, 120, 120)
        set_cell_margins(cell_d, 100, 100, 120, 120)
        
        p1 = cell_t.paragraphs[0]
        r1 = p1.add_run(title)
        r1.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(9.5)
        
        p2 = cell_d.paragraphs[0]
        r2 = p2.add_run(desc)
        r2.font.name = 'Arial'
        r2.font.size = Pt(9.5)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # --- PHẦN 5: CẬP NHẬT PHIÊN BẢN ---
    h5 = doc.add_heading(level=1)
    h5.paragraph_format.space_before = Pt(14)
    h5.paragraph_format.space_after = Pt(6)
    r = h5.add_run("5. Cập Nhật Khi Có Phiên Bản Mới (OTA Update)")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "Khi GO4AI có bản nâng cấp tính năng hoặc thêm mẫu template mới:\n"
        "1. Nhấp đúp vào file '2-Cap-Nhat-Windows.bat'.\n"
        "2. Hệ thống sẽ tự động cập nhật bản mới nhất từ GitHub trong 30–60 giây.\n"
        "3. Toàn bộ kịch bản, dự án và video bạn đã làm đều được GIỮ NGUYÊN 100%."
    ).font.name = 'Arial'
    
    # --- PHẦN 6: XỬ LÝ SỰ CỐ ---
    h6 = doc.add_heading(level=1)
    h6.paragraph_format.space_before = Pt(14)
    h6.paragraph_format.space_after = Pt(6)
    r = h6.add_run("6. Bảng Xử Lý Sự Cố Thường Gặp Trên Windows")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    troubles = [
        ("Tình huống 1: Windows hiện bảng xanh 'Windows protected your PC'",
         "Đây là tính năng SmartScreen của Windows khi mở file .bat mới tải từ mạng về.\n"
         "👉 Cách xử lý: Bấm vào dòng chữ 'More info' ➔ Bấm nút 'Run anyway' là xong."),
        
        ("Tình huống 2: Lỗi tải Node.js tự động không thành công do mạng chậm",
         "Nếu mạng chặn tải tự động:\n"
         "👉 Cách xử lý: Truy cập https://nodejs.org/en/download ➔ Chọn 'Windows Installer (.msi)' ➔ Tải về và cài đặt bình thường ➔ Mở lại file 1-Chay-Tren-Windows.bat."),
         
        ("Tình huống 3: Báo lỗi cổng 'Port 3075 already in use'",
         "Do có một phiên bản app trước đó đang chạy ngầm trong máy.\n"
         "👉 Cách xử lý: Nhấn tổ hợp phím Ctrl + Shift + Esc để mở Task Manager ➔ Tìm tiến trình 'Node.js: Server-side JavaScript' (hoặc node.exe) ➔ Click chuột phải chọn 'End task' ➔ Mở lại file .bat."),
         
        ("Tình huống 4: Cài đặt lần đầu tải lâu quá",
         "Quá trình lần đầu cần tải khoảng 200–300MB thư viện cần thiết. Hãy đảm bảo máy tính kết nối Wifi ổn định và không tắt cửa sổ đen trong lúc cài đặt.")
    ]
    
    for title, fix in troubles:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(6)
        p_t.paragraph_format.space_after = Pt(2)
        r_t = p_t.add_run(f"🔸 {title}")
        r_t.bold = True
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = RGBColor(185, 28, 28)
        
        p_f = doc.add_paragraph()
        p_f.paragraph_format.space_before = Pt(2)
        p_f.paragraph_format.space_after = Pt(4)
        r_f = p_f.add_run(fix)
        r_f.font.name = 'Arial'
        r_f.font.size = Pt(9.5)
        r_f.font.color.rgb = RGBColor(51, 65, 85)
        
    # --- PHẦN 7: LIÊN HỆ ---
    h7 = doc.add_heading(level=1)
    h7.paragraph_format.space_before = Pt(14)
    h7.paragraph_format.space_after = Pt(6)
    r = h7.add_run("7. Hỗ Trợ Kỹ Thuật & Cộng Đồng")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Nếu bạn cần trợ giúp thêm, hãy liên hệ ngay với đội ngũ kỹ thuật GO4AI:\n\n"
        "• 📧 Email tiếp nhận hỗ trợ: hocvien@go4ai.life (Phản hồi nhanh nhất)\n"
        "• 💬 Tính năng hỗ trợ trực tiếp: Bấm nút 'Góp ý & Báo lỗi' ở góc trên thanh công cụ trong app\n"
        "• 🌐 Website cộng đồng: https://go4ai.life\n\n"
        "Chúc bạn có những trải nghiệm sáng tạo video AI tuyệt vời cùng GO4AI Studio!"
    ).font.name = 'Arial'
    
    out_path = os.path.join(os.getcwd(), "HUONG-DAN-CAI-DAT-WINDOWS.docx")
    doc.save(out_path)
    print(f"[OK] Da tao thanh cong file DOCX: {out_path}")

if __name__ == "__main__":
    build_windows_docx()
