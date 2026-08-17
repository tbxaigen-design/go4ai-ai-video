import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_callout_box(doc, text_list, title="💡 LƯU Ý QUAN TRỌNG", bg_hex="F0F9FF", border_color="0284C7"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=220)
    
    # Border
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:top w:val="none"/>
            <w:right w:val="none"/>
            <w:bottom w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    # Title
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(title)
    run_t.bold = True
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(15, 23, 42)
    
    # Content
    for line in text_list:
        p_line = cell.add_paragraph()
        p_line.paragraph_format.space_before = Pt(2)
        p_line.paragraph_format.space_after = Pt(3)
        run = p_line.add_run(line)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 41, 59)
    
    # Space after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def build_docx():
    doc = Document()
    
    # Page setup (A4, 0.75 inch margins)
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Document Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("🍏 GO4AI AI VIDEO STUDIO")
    run_title.bold = True
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(30, 64, 175) # #1E40AF
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("Cẩm Nang Hướng Dẫn Cài Đặt & Sử Dụng Chi Tiết Cho MacBook (macOS)\n(Dành cho người mới bắt đầu / No-code)")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(71, 85, 105)
    run_sub.italic = True
    
    # Metadata Badge Table
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cols = [
        ("📱 Nền tảng", "macOS 12+ (M1/M2/M3/M4 & Intel)"),
        ("⚡ Phiên bản", "Beta 1.0 (Local Render)"),
        ("🤝 Hỗ trợ", "hocvien@go4ai.life")
    ]
    for i, (k, v) in enumerate(cols):
        cell = meta_table.cell(0, i)
        cell.width = Inches(2.15)
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, 80, 80, 100, 100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{k}\n")
        r1.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(8.5)
        r1.font.color.rgb = RGBColor(30, 41, 59)
        r2 = p.add_run(v)
        r2.font.name = 'Arial'
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # --- PHẦN 1: GIỚI THIỆU ---
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("1. Giới Thiệu & Nguyên Lý Hoạt Động")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "GO4AI AI Video Studio là ứng dụng tạo và biên tập video tự động bằng trí tuệ nhân tạo. "
        "Điểm đặc biệt là ứng dụng xử lý và xuất video MP4 hoàn toàn trực tiếp trên máy tính của bạn (Local Render), "
        "giúp bảo mật dữ liệu tuyệt đối và xuất video Full HD 60fps cực nhanh nhờ tận dụng sức mạnh phần cứng của máy Mac."
    ).font.name = 'Arial'
    
    add_callout_box(
        doc,
        [
            "• Bạn KHÔNG cần biết lập trình, KHÔNG cần biết gõ code phức tạp.",
            "• Tất cả thao tác cài đặt đều được tự động hóa chỉ bằng 1 cú nhấp chuột.",
            "• Nếu thấy xuất hiện một cửa sổ màu đen (Terminal), đừng lo lắng! Đó là ứng dụng đang tự động chuẩn bị môi trường chạy ngầm."
        ],
        title="🌟 DÀNH CHO BẠN (NGƯỜI DÙNG NO-CODE)",
        bg_hex="ECFDF5",
        border_color="059669"
    )
    
    # --- PHẦN 2: CHUẨN BỊ ---
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)
    r = h2.add_run("2. Chuẩn Bị Trước Khi Cài Đặt (Chỉ 1 bước)")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Để ứng dụng có thể chạy được trên MacBook, máy tính cần có môi trường Node.js (phiên bản 20 trở lên). "
        "Nếu máy bạn chưa từng cài đặt, hãy thực hiện cài đặt như sau:"
    ).font.name = 'Arial'
    
    steps_node = [
        ("Bước 2.1: Tải bộ cài Node.js", "Truy cập website chính thức: https://nodejs.org/en/download và chọn bản LTS (Recommended For Most Users). File tải về sẽ có dạng .pkg (ví dụ: node-v20.x.x.pkg)."),
        ("Bước 2.2: Cài đặt vào máy", "Mở file .pkg vừa tải về trong thư mục Downloads. Bấm 'Tiếp tục' (Continue) ➔ 'Đồng ý' (Agree) ➔ 'Cài đặt' (Install) ➔ Nhập mật khẩu mở máy Mac nếu được hỏi."),
        ("Bước 2.3: Hoàn tất", "Bấm 'Đóng' (Close) khi màn hình báo cài đặt thành công. Bây giờ máy bạn đã sẵn sàng 100%!")
    ]
    
    for title, desc in steps_node:
        p_step = doc.add_paragraph()
        p_step.paragraph_format.space_before = Pt(3)
        p_step.paragraph_format.space_after = Pt(3)
        r_t = p_step.add_run(f"👉 {title}: ")
        r_t.bold = True
        r_t.font.name = 'Arial'
        r_t.font.color.rgb = RGBColor(15, 23, 42)
        r_d = p_step.add_run(desc)
        r_d.font.name = 'Arial'
        r_d.font.color.rgb = RGBColor(51, 65, 85)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # --- PHẦN 3: CÁC BƯỚC CÀI ĐẶT ---
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    r = h3.add_run("3. Hướng Dẫn Cài Đặt & Mở Ứng Dụng (3 Bước)")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    # Bước 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("📥 BƯỚC 1: Tải bộ phần mềm về máy")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(2, 132, 199)
    
    p_desc1 = doc.add_paragraph()
    p_desc1.paragraph_format.space_after = Pt(4)
    p_desc1.add_run(
        "1. Truy cập đường link tải: https://github.com/tbxaigen-design/go4ai-ai-video\n"
        "2. Nhìn sang góc phải màn hình, bấm vào nút màu xanh lá cây có chữ '< > Code'.\n"
        "3. Trong bảng chọn hiện ra, bấm vào dòng chữ 'Download ZIP'.\n"
        "4. File nén 'go4ai-ai-video-main.zip' sẽ tự động tải về thư mục Downloads của bạn.\n"
        "5. Nhấp đúp vào file ZIP để giải nén thành thư mục 'go4ai-ai-video-main'."
    ).font.name = 'Arial'
    
    # Bước 2
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("🚀 BƯỚC 2: Mở file khởi động lần đầu")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(2, 132, 199)
    
    p_desc2 = doc.add_paragraph()
    p_desc2.paragraph_format.space_after = Pt(4)
    p_desc2.add_run(
        "Mở thư mục vừa giải nén, bạn sẽ thấy file có tên: 1-Chay-Tren-Macbook.command\n"
        "Do chính sách bảo mật của Apple, trong LẦN ĐẦU TIÊN mở ứng dụng, bạn thực hiện như sau:\n"
        "1. Click chuột phải (hoặc giữ phím Control + Click chuột) vào file '1-Chay-Tren-Macbook.command'.\n"
        "2. Chọn 'Mở' (Open) từ menu danh sách.\n"
        "3. Nếu Mac hiện bảng hỏi xác nhận, bấm nút 'Mở' (Open) một lần nữa."
    ).font.name = 'Arial'
    
    add_callout_box(
        doc,
        [
            "• Một cửa sổ dòng lệnh màu đen (Terminal) sẽ mở ra và tự động thực hiện mọi thứ.",
            "• Hệ thống sẽ tự động cài đặt thư viện, bộ nạp video FFmpeg, trình duyệt Chromium và giọng đọc AI.",
            "• Quá trình này mất khoảng 2–5 phút tuỳ tốc độ mạng Internet. Vui lòng KHÔNG TẮT cửa sổ màu đen này!"
        ],
        title="⏳ QUÁ TRÌNH TỰ ĐỘNG CHUẨN BỊ (2-5 PHÚT)",
        bg_hex="FFFBEB",
        border_color="D97706"
    )
    
    # Bước 3
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("🌐 BƯỚC 3: Trải nghiệm GO4AI Studio trên trình duyệt")
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(2, 132, 199)
    
    p_desc3 = doc.add_paragraph()
    p_desc3.paragraph_format.space_after = Pt(6)
    p_desc3.add_run(
        "Sau khi chuẩn bị xong, trình duyệt web (Safari/Chrome) trên máy Mac sẽ tự động mở trang web: http://127.0.0.1:3075\n"
        "Tại đây, bạn đã có thể bắt đầu chọn template, nhập kịch bản, chỉnh sửa hiệu ứng và xuất video chất lượng cao!"
    ).font.name = 'Arial'
    
    # --- PHẦN 4: SỬ DỤNG HÀNG NGÀY ---
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(6)
    r = h4.add_run("4. Hướng Dẫn Sử Dụng Hàng Ngày & Tắt Ứng Dụng")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    usage_table = doc.add_table(rows=2, cols=2)
    usage_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    u_data = [
        ("🟢 Cách Mở App Hàng Ngày", "Từ lần thứ 2 trở đi, bạn chỉ cần nhấp đúp (Double click) vào file '1-Chay-Tren-Macbook.command'. Ứng dụng sẽ khởi động ngay lập tức trong 2-3 giây."),
        ("🔴 Cách Tắt App Đúng Cách", "Khi làm việc xong, bạn quay lại cửa sổ Terminal màu đen, bấm tổ hợp phím Ctrl + C trên bàn phím rồi đóng cửa sổ. Không tắt đột ngột khi đang render video.")
    ]
    for row_idx, (title, desc) in enumerate(u_data):
        cell_t = usage_table.cell(row_idx, 0)
        cell_d = usage_table.cell(row_idx, 1)
        cell_t.width = Inches(2.3)
        cell_d.width = Inches(4.2)
        set_cell_background(cell_t, "F8FAFC")
        set_cell_background(cell_d, "FFFFFF")
        set_cell_margins(cell_t, 100, 100, 120, 120)
        set_cell_margins(cell_d, 100, 100, 120, 120)
        
        p1 = cell_t.paragraphs[0]
        r1 = p1.add_run(title)
        r1.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(9.5)
        
        p2 = cell_d.paragraphs[0]
        r2 = p2.add_run(desc)
        r2.font.name = 'Arial'
        r2.font.size = Pt(9.5)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # --- PHẦN 5: CẬP NHẬT PHIÊN BẢN ---
    h5 = doc.add_heading(level=1)
    h5.paragraph_format.space_before = Pt(14)
    h5.paragraph_format.space_after = Pt(6)
    r = h5.add_run("5. Cập Nhật Khi Có Phiên Bản Mới (OTA Update)")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(
        "Khi đội ngũ GO4AI phát hành tính năng mới hoặc mẫu video mới, bạn không cần phải tải lại từ đầu:\n"
        "1. Click chuột phải vào file '2-Cap-Nhat-Macbook.command' ➔ Chọn 'Mở' (Open).\n"
        "2. Hệ thống sẽ tự động đồng bộ bản mới nhất từ GitHub trong khoảng 30–60 giây.\n"
        "3. Toàn bộ kịch bản, video và dự án bạn đã tạo đều được BẢO TOÀN 100% (không bị mất)."
    ).font.name = 'Arial'
    
    # --- PHẦN 6: XỬ LÝ SỰ CỐ ---
    h6 = doc.add_heading(level=1)
    h6.paragraph_format.space_before = Pt(14)
    h6.paragraph_format.space_after = Pt(6)
    r = h6.add_run("6. Cẩm Nang Xử Lý Sự Cố Thường Gặp Trên Mac")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    troubles = [
        ("Tình huống 1: Mac báo 'Nhà phát triển chưa xác minh' (Unverified Developer)",
         "Cách 1: Luôn nhớ Click chuột phải vào file .command ➔ Chọn 'Open' (Mở) thay vì click đúp.\n"
         "Cách 2: Vào menu Apple  ở góc trên trái ➔ 'Cài đặt hệ thống' (System Settings) ➔ 'Quyền riêng tư & Bảo mật' (Privacy & Security) ➔ Cuộn xuống mục Bảo mật ➔ Bấm 'Vẫn mở' (Open Anyway)."),
        
        ("Tình huống 2: Mac báo 'File bị hỏng và không thể mở' (File is damaged)",
         "Do tính năng bảo vệ file tải về từ Internet của Apple. Bạn chỉ cần:\n"
         "1. Nhấn tổ hợp phím Cmd + Space để mở Spotlight ➔ Gõ 'Terminal' rồi nhấn Enter.\n"
         "2. Copy và dán lệnh sau vào cửa sổ Terminal rồi bấm Enter:\n"
         "   xattr -cr ~/Downloads/go4ai-ai-video-main\n"
         "3. Sau đó mở lại file .command bình thường."),
         
        ("Tình huống 3: Báo lỗi 'Permission denied' khi mở file",
         "1. Mở ứng dụng Terminal trên Mac.\n"
         "2. Gõ chữ 'chmod +x ' (có dấu cách ở cuối), sau đó kéo file '1-Chay-Tren-Macbook.command' thả vào cửa sổ Terminal rồi bấm Enter."),
         
        ("Tình huống 4: Báo lỗi 'Port 3075 already in use'",
         "Do có một phiên bản cũ đang chạy ngầm trong máy:\n"
         "1. Mở Terminal, gõ lệnh: pkill node rồi bấm Enter.\n"
         "2. Khởi động lại file '1-Chay-Tren-Macbook.command'.")
    ]
    
    for title, fix in troubles:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(6)
        p_t.paragraph_format.space_after = Pt(2)
        r_t = p_t.add_run(f"🔸 {title}")
        r_t.bold = True
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(10)
        r_t.font.color.rgb = RGBColor(185, 28, 28) # red
        
        p_f = doc.add_paragraph()
        p_f.paragraph_format.space_before = Pt(2)
        p_f.paragraph_format.space_after = Pt(4)
        r_f = p_f.add_run(fix)
        r_f.font.name = 'Arial'
        r_f.font.size = Pt(9.5)
        r_f.font.color.rgb = RGBColor(51, 65, 85)
        
    # --- PHẦN 7: LIÊN HỆ ---
    h7 = doc.add_heading(level=1)
    h7.paragraph_format.space_before = Pt(14)
    h7.paragraph_format.space_after = Pt(6)
    r = h7.add_run("7. Hỗ Trợ Kỹ Thuật & Cộng Đồng")
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(30, 64, 175)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Nếu bạn gặp bất kỳ khó khăn nào trong quá trình cài đặt và trải nghiệm, hãy liên hệ ngay với đội ngũ kỹ thuật GO4AI:\n\n"
        "• 📧 Email tiếp nhận hỗ trợ: hocvien@go4ai.life (Phản hồi nhanh nhất)\n"
        "• 💬 Tính năng hỗ trợ trực tiếp: Bấm nút 'Góp ý & Báo lỗi' ở góc trên thanh công cụ trong app\n"
        "• 🌐 Website cộng đồng: https://go4ai.life\n\n"
        "Chúc bạn có những trải nghiệm sáng tạo video AI tuyệt vời cùng GO4AI Studio!"
    ).font.name = 'Arial'
    
    out_path = os.path.join(os.getcwd(), "HUONG-DAN-CAI-DAT-MACBOOK.docx")
    doc.save(out_path)
    print(f"[OK] Da tao thanh cong file DOCX: {out_path}")

if __name__ == "__main__":
    build_docx()
